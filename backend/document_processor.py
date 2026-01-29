"""Document processing pipeline for text extraction and chunking."""

import uuid
from pathlib import Path
from typing import List, Tuple
import re

# PDF processing
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

# DOCX processing
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

from models import DocumentChunk, ChunkMetadata
from config import settings


class DocumentProcessor:
    """Handles document text extraction and chunking."""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt'}
    
    def __init__(self, chunk_size: int = None, chunk_overlap: int = None):
        """
        Initialize document processor.
        
        Args:
            chunk_size: Size of text chunks (default from settings)
            chunk_overlap: Overlap between chunks (default from settings)
        """
        self.chunk_size = chunk_size or settings.chunk_size
        self.chunk_overlap = chunk_overlap or settings.chunk_overlap
    
    def validate_file(self, filename: str) -> Tuple[bool, str]:
        """
        Validate if file type is supported.
        
        Args:
            filename: Name of the file
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        file_path = Path(filename)
        extension = file_path.suffix.lower()
        
        if extension not in self.SUPPORTED_EXTENSIONS:
            return False, f"Unsupported file type: {extension}. Supported: {', '.join(self.SUPPORTED_EXTENSIONS)}"
        
        if extension == '.pdf' and PdfReader is None:
            return False, "PDF support not installed. Install pypdf: pip install pypdf"
        
        if extension == '.docx' and DocxDocument is None:
            return False, "DOCX support not installed. Install python-docx: pip install python-docx"
        
        return True, ""
    
    def extract_text(self, file_path: str) -> Tuple[str, dict]:
        """
        Extract text from a document file.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension == '.pdf':
            return self._extract_pdf(file_path)
        elif extension == '.docx':
            return self._extract_docx(file_path)
        elif extension == '.txt':
            return self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    def _extract_pdf(self, file_path: str) -> Tuple[str, dict]:
        """Extract text from PDF file."""
        reader = PdfReader(file_path)
        text_parts = []
        page_map = {}  # Track which text came from which page
        
        current_pos = 0
        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            text_parts.append(page_text)
            page_map[current_pos] = page_num
            current_pos += len(page_text)
        
        full_text = "\n\n".join(text_parts)
        metadata = {
            'num_pages': len(reader.pages),
            'page_map': page_map
        }
        
        return full_text, metadata
    
    def _extract_docx(self, file_path: str) -> Tuple[str, dict]:
        """Extract text from DOCX file."""
        doc = DocxDocument(file_path)
        text_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        full_text = "\n\n".join(text_parts)
        
        metadata = {
            'num_paragraphs': len(doc.paragraphs)
        }
        
        return full_text, metadata
    
    def _extract_txt(self, file_path: str) -> Tuple[str, dict]:
        """Extract text from TXT file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        metadata = {
            'encoding': 'utf-8'
        }
        
        return text, metadata
    
    def chunk_text(self, text: str, source: str, page_map: dict = None) -> List[DocumentChunk]:
        """
        Split text into overlapping chunks.
        
        Args:
            text: Text to chunk
            source: Source filename
            page_map: Optional mapping of text positions to page numbers
            
        Returns:
            List of DocumentChunk objects
        """
        # Clean text
        text = self._clean_text(text)
        
        if len(text) == 0:
            raise ValueError("Extracted text is empty")
        
        chunks = []
        start = 0
        
        while start < len(text):
            # Calculate end position
            end = start + self.chunk_size
            
            # If not at the end, try to break at a sentence or word boundary
            if end < len(text):
                # Look for sentence boundary (. ! ?)
                sentence_end = max(
                    text.rfind('. ', start, end),
                    text.rfind('! ', start, end),
                    text.rfind('? ', start, end)
                )
                
                if sentence_end > start:
                    end = sentence_end + 1
                else:
                    # Fall back to word boundary
                    space_pos = text.rfind(' ', start, end)
                    if space_pos > start:
                        end = space_pos
            
            # Extract chunk
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                # Determine page number if page_map provided
                page = None
                if page_map:
                    for pos, page_num in sorted(page_map.items(), reverse=True):
                        if start >= pos:
                            page = page_num
                            break
                
                chunk_id = str(uuid.uuid4())
                metadata = ChunkMetadata(
                    source=source,
                    page=page,
                    chunk_id=chunk_id
                )
                
                chunk = DocumentChunk(
                    chunk_id=chunk_id,
                    text=chunk_text,
                    metadata=metadata
                )
                chunks.append(chunk)
            
            # Move to next chunk with overlap
            start = end - self.chunk_overlap
            
            # Ensure we make progress
            if start <= chunks[-1].text.find(chunks[-1].text[:50]) if chunks else 0:
                start = end
        
        return chunks
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove excessive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()
    
    def process_document(self, file_path: str) -> Tuple[List[DocumentChunk], dict]:
        """
        Complete pipeline: extract and chunk document.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Tuple of (chunks, extraction_metadata)
        """
        # Validate
        filename = Path(file_path).name
        is_valid, error = self.validate_file(filename)
        if not is_valid:
            raise ValueError(error)
        
        # Extract
        text, metadata = self.extract_text(file_path)
        
        # Validate extraction
        if len(text.strip()) < 10:
            raise ValueError(f"Extracted text too short ({len(text)} chars). Possible OCR or extraction issue.")
        
        # Chunk
        page_map = metadata.get('page_map')
        chunks = self.chunk_text(text, filename, page_map)
        
        return chunks, metadata
